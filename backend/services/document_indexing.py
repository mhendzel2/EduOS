from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Optional

from storage.vector_store import Document

logger = logging.getLogger(__name__)

_CHUNK_SIZE = 1500
_CHUNK_OVERLAP = 200
_BINARY_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".webp",
    ".bmp",
    ".ico",
    ".mp4",
    ".mov",
    ".webm",
    ".m4v",
    ".avi",
    ".mp3",
    ".wav",
}

# ---------------------------------------------------------------------------
# Section detection
# ---------------------------------------------------------------------------

# Standard biology paper section headers — order matters (matched top-down)
_SECTION_PATTERNS: list[tuple[str, re.Pattern]] = [
    ("abstract",      re.compile(r"^\s*abstract\s*$",                         re.IGNORECASE | re.MULTILINE)),
    ("introduction",  re.compile(r"^\s*introduction\s*$",                     re.IGNORECASE | re.MULTILINE)),
    ("results",       re.compile(r"^\s*results?\s*$",                         re.IGNORECASE | re.MULTILINE)),
    ("methods",       re.compile(r"^\s*(?:materials?\s+(?:and\s+)?)?methods?\s*$",  re.IGNORECASE | re.MULTILINE)),
    ("discussion",    re.compile(r"^\s*discussion\s*$",                       re.IGNORECASE | re.MULTILINE)),
    ("conclusion",    re.compile(r"^\s*conclusions?\s*$",                     re.IGNORECASE | re.MULTILINE)),
    ("references",    re.compile(r"^\s*references?\s*$",                      re.IGNORECASE | re.MULTILINE)),
    ("supplementary", re.compile(r"^\s*supplementar(?:y|ies)\s",              re.IGNORECASE | re.MULTILINE)),
]

# Figure/table reference capture
_RE_FIGURE_REF = re.compile(
    r"\b((?:Fig(?:ure)?|Suppl?\.?\s*Fig(?:ure)?|Extended\s+Data\s+Fig(?:ure)?|"
    r"Table|Suppl?\.?\s*Table)\s*\d+[A-Za-z]?(?:[-–]\d+[A-Za-z]?)?)\b",
    re.IGNORECASE,
)


def _detect_sections(text: str) -> dict[str, str]:
    """Return a dict mapping section_label -> section_text.

    If fewer than 2 sections are detected, returns {"body": text} as a fallback
    with a "low_confidence" key set to True (accessible via the dict itself).
    """
    # Find all section header positions
    hits: list[tuple[int, str]] = []
    for label, pattern in _SECTION_PATTERNS:
        for m in pattern.finditer(text):
            hits.append((m.start(), label))

    # Sort by position; deduplicate overlapping labels
    hits.sort(key=lambda x: x[0])
    seen_labels: set[str] = set()
    unique_hits: list[tuple[int, str]] = []
    for pos, label in hits:
        if label not in seen_labels:
            seen_labels.add(label)
            unique_hits.append((pos, label))

    if len(unique_hits) < 2:
        return {"body": text, "_low_confidence": True}  # type: ignore[return-value]

    sections: dict[str, str] = {}
    for i, (start, label) in enumerate(unique_hits):
        end = unique_hits[i + 1][0] if i + 1 < len(unique_hits) else len(text)
        sections[label] = text[start:end].strip()

    return sections


def _extract_figure_refs(text: str) -> list[str]:
    seen: set[str] = set()
    refs: list[str] = []
    for m in _RE_FIGURE_REF.finditer(text):
        val = m.group(1).strip()
        key = val.lower()
        if key not in seen:
            seen.add(key)
            refs.append(val)
    return refs


# ---------------------------------------------------------------------------
# Flat chunker (unchanged, used as fallback)
# ---------------------------------------------------------------------------

def _chunk_documents(text: str, metadata: dict, *, id_prefix: Optional[str] = None) -> list[Document]:
    chunks: list[Document] = []
    start = 0
    chunk_index = 0
    while start < len(text):
        end = start + _CHUNK_SIZE
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunk_id = f"{id_prefix}:chunk:{chunk_index}" if id_prefix else None
            chunks.append(
                Document(
                    id=chunk_id,
                    content=chunk_text,
                    metadata={**metadata, "chunk_index": chunk_index},
                )
            )
            chunk_index += 1
        start += _CHUNK_SIZE - _CHUNK_OVERLAP
    return chunks


# ---------------------------------------------------------------------------
# Section-aware chunker
# ---------------------------------------------------------------------------

def section_aware_chunk_documents(
    text: str,
    metadata: dict,
    *,
    id_prefix: Optional[str] = None,
) -> list[Document]:
    """Chunk *text* with section awareness, nomenclature resolution, and
    methodological fingerprinting on the Methods section.

    Each chunk's metadata includes:
        - section: which paper section the chunk came from
        - figure_refs: pipe-delimited figure/table references found in the chunk
        - mf_* keys from MethodologicalFingerprint (methods section only,
          then propagated to all chunks from the same document)
        - mf_rigor_score, mf_rigor_flags (always present)
    """
    # Lazy imports to avoid startup overhead
    try:
        from utils.nomenclature_resolver import resolve_text
    except ImportError:
        def resolve_text(t: str) -> str:  # type: ignore[misc]
            return t

    try:
        from services.methodological_fingerprinting import fingerprint_text, fingerprint_empty
    except ImportError:
        def fingerprint_text(t: str, **_):  # type: ignore[misc]
            from services.methodological_fingerprinting import fingerprint_empty
            return fingerprint_empty()
        def fingerprint_empty():  # type: ignore[misc]
            class _FP:
                rigor_score = 0.0
                rigor_flags: list = []
                def to_metadata(self): return {"mf_rigor_score": 0.0, "mf_rigor_flags": ""}
            return _FP()

    sections = _detect_sections(text)
    low_confidence: bool = bool(sections.pop("_low_confidence", False))

    if low_confidence:
        logger.debug("section_aware_chunk_documents: fewer than 2 sections detected, falling back to flat chunking")
        resolved = resolve_text(text)
        fp = fingerprint_text(resolved)
        fp_meta = fp.to_metadata()
        base = {**metadata, **fp_meta, "section": "body", "low_confidence_sections": True}
        return _chunk_documents(resolved, base, id_prefix=id_prefix)

    # Run fingerprinting on the methods section if present
    methods_text = sections.get("methods", "")
    if methods_text:
        fp = fingerprint_text(methods_text, methods_only=True)
    else:
        fp = fingerprint_empty()
    fp_meta = fp.to_metadata()

    chunks: list[Document] = []
    global_chunk_index = 0

    for section_label, section_text in sections.items():
        resolved_section = resolve_text(section_text)
        section_figure_refs = _extract_figure_refs(resolved_section)

        # Chunk within the section
        start = 0
        while start < len(resolved_section):
            end = start + _CHUNK_SIZE
            chunk_text = resolved_section[start:end].strip()
            if chunk_text:
                chunk_figure_refs = _extract_figure_refs(chunk_text)
                chunk_id = f"{id_prefix}:chunk:{global_chunk_index}" if id_prefix else None
                chunk_meta = {
                    **metadata,
                    **fp_meta,
                    "section": section_label,
                    "chunk_index": global_chunk_index,
                    "figure_refs": "|".join(chunk_figure_refs) if chunk_figure_refs else "",
                    "section_figure_refs": "|".join(section_figure_refs) if section_figure_refs else "",
                    "low_confidence_sections": False,
                }
                chunks.append(Document(id=chunk_id, content=chunk_text, metadata=chunk_meta))
                global_chunk_index += 1
            start += _CHUNK_SIZE - _CHUNK_OVERLAP

    if not chunks:
        # Last-resort fallback
        resolved = resolve_text(text)
        return _chunk_documents(resolved, {**metadata, **fp_meta}, id_prefix=id_prefix)

    return chunks


# ---------------------------------------------------------------------------
# Public entry points
# ---------------------------------------------------------------------------

async def build_vector_documents_for_file(
    path: str,
    *,
    document_id: str,
    base_metadata: dict,
) -> list[Document]:
    text_content = await extract_text_for_file(path)
    if not text_content.strip():
        return []

    metadata = {**base_metadata, "document_id": document_id, "source_file": Path(path).name}
    chunked = section_aware_chunk_documents(text_content, metadata, id_prefix=document_id)
    if chunked:
        return chunked

    return [Document(id=document_id, content=text_content, metadata=metadata)]


def build_vector_documents_for_structured_document(
    payload: dict,
    *,
    document_id: str,
    base_metadata: dict,
) -> list[Document]:
    title = str(payload.get("title") or "").strip()
    abstract = str(payload.get("abstract") or "").strip()
    content = str(payload.get("content") or "").strip()
    citation = str(payload.get("citation") or "").strip()
    source_type = str(payload.get("source_type") or "structured").strip()
    source_identifier = str(payload.get("source_identifier") or "").strip()
    source_url = str(payload.get("source_url") or "").strip()
    authors = [str(item).strip() for item in list(payload.get("authors") or []) if str(item).strip()]

    sections = [
        section
        for section in [
            f"Title: {title}" if title else "",
            f"Abstract: {abstract}" if abstract else "",
            f"Content: {content}" if content else "",
            f"Authors: {', '.join(authors)}" if authors else "",
            f"Citation: {citation}" if citation else "",
            f"Source Type: {source_type}" if source_type else "",
            f"Identifier: {source_identifier}" if source_identifier else "",
            f"Source URL: {source_url}" if source_url else "",
        ]
        if section
    ]
    normalized_text = "\n\n".join(sections).strip()
    if not normalized_text:
        return []

    metadata = {
        **base_metadata,
        "document_id": document_id,
        "source_type": source_type,
        "source_identifier": source_identifier,
        "source_url": source_url,
        "citation": citation,
        "authors": authors,
    }
    chunked = section_aware_chunk_documents(normalized_text, metadata, id_prefix=document_id)
    if chunked:
        return chunked
    return [Document(id=document_id, content=normalized_text, metadata=metadata)]


async def extract_text_for_file(path: str) -> str:
    suffix = Path(path).suffix.lower()
    if suffix in _BINARY_EXTENSIONS:
        return ""
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".json":
        return _extract_json_text(path)
    return _read_text_file(path)


def _read_text_file(path: str) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    except Exception:
        logger.warning("Failed to read %s as plain text", path, exc_info=True)
        return ""


def _extract_pdf_text(path: str) -> str:
    try:
        import fitz

        with fitz.open(path) as document:
            return "\n\n".join(page.get_text("text") for page in document).strip()
    except Exception:
        logger.warning("Failed to extract PDF text from %s", path, exc_info=True)
        return ""


def _extract_docx_text(path: str) -> str:
    try:
        from docx import Document as DocxDocument

        document = DocxDocument(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        logger.warning("Failed to extract DOCX text from %s", path, exc_info=True)
        return ""


def _extract_json_text(path: str) -> str:
    try:
        parsed = json.loads(Path(path).read_text(encoding="utf-8", errors="ignore"))
        return json.dumps(parsed, indent=2, ensure_ascii=True)
    except Exception:
        logger.warning("Failed to normalize JSON text from %s", path, exc_info=True)
        return _read_text_file(path)
